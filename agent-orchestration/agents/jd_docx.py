"""
Docx builder for JD Builder module.

build_docx(doc, brand_color, logo_bytes) -> bytes
    Renders a JDDocument as a branded .docx file.

extract_jd_text(doc) -> str
    Flattens all prose fields for Bedrock guardrail OUTPUT check.
"""

import io
from typing import Optional

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm, Pt, RGBColor

from models.context import JDDocument


_DEFAULT_COLOR = "#1D3557"


def _hex_to_rgb(hex_color: str) -> tuple[int, int, int]:
    """Parse a hex color string (with or without #) to an (r, g, b) tuple."""
    hex_color = hex_color.lstrip("#").strip()
    try:
        r = int(hex_color[0:2], 16)
        g = int(hex_color[2:4], 16)
        b = int(hex_color[4:6], 16)
        return r, g, b
    except (ValueError, IndexError):
        return _hex_to_rgb(_DEFAULT_COLOR)


def _rgb_color(hex_color: str) -> RGBColor:
    r, g, b = _hex_to_rgb(hex_color)
    return RGBColor(r, g, b)


def _add_heading(doc: Document, text: str, brand_color: str, level: int = 1) -> None:
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = _rgb_color(brand_color)
        run.font.bold = True
        run.font.size = Pt(13 if level == 1 else 11)


def _add_bullet_list(doc: Document, items: list[str]) -> None:
    for item in items:
        if item:
            doc.add_paragraph(item, style="List Bullet")


def _add_accent_bar(doc: Document, brand_color: str) -> None:
    """Add a thin colored horizontal rule using a bottom border on an empty paragraph."""
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    r, g, b = _hex_to_rgb(brand_color)
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "6")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), f"{r:02X}{g:02X}{b:02X}")
    pBdr.append(bottom)
    pPr.append(pBdr)


def _set_header_footer(doc: Document, role_title: str, brand_color: str) -> None:
    from datetime import date

    section = doc.sections[0]

    # Header — role title left, date right
    header = section.header
    header.is_linked_to_previous = False
    page_width = section.page_width - section.left_margin - section.right_margin
    htable = header.add_table(1, 2, width=page_width)
    htable.style = "Table Grid"
    htable.rows[0].cells[0].text = role_title
    htable.rows[0].cells[1].text = date.today().strftime("%d %b %Y")
    htable.rows[0].cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for cell in htable.rows[0].cells:
        for para in cell.paragraphs:
            for run in para.runs:
                run.font.color.rgb = _rgb_color(brand_color)
                run.font.size = Pt(9)
    # Remove table borders
    for cell in htable.rows[0].cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = OxmlElement("w:tcBorders")
        for border_name in ("top", "left", "bottom", "right"):
            border = OxmlElement(f"w:{border_name}")
            border.set(qn("w:val"), "none")
            tcBorders.append(border)
        tcPr.append(tcBorders)

    # Footer
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    fp.text = "Built with JDQI  ·  Confidential"
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in fp.runs:
        run.font.size = Pt(8)
        run.font.color.rgb = RGBColor(0x80, 0x80, 0x80)


def _add_logo(doc: Document, logo_bytes: Optional[bytes]) -> None:
    if logo_bytes:
        stream = io.BytesIO(logo_bytes)
        try:
            doc.add_picture(stream, width=Cm(3))
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.RIGHT
            return
        except Exception:
            pass
    # Placeholder
    p = doc.add_paragraph("[Company Logo]")
    p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in p.runs:
        run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)
        run.font.size = Pt(9)


def build_docx(
    doc: JDDocument,
    brand_color: str,
    logo_bytes: Optional[bytes],
) -> bytes:
    """
    Render a JDDocument as a branded .docx file.

    Args:
        doc:         JDDocument dict from the Draft agent.
        brand_color: Hex color string (e.g. "#1D3557") for headings and accents.
        logo_bytes:  Raw image bytes (PNG/JPG) or None for placeholder.

    Returns:
        Raw .docx bytes suitable for st.download_button.
    """
    d = Document()

    style = d.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(11)
    style.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)

    _set_header_footer(d, doc.get("role_title", "Job Description"), brand_color)
    _add_logo(d, logo_bytes)
    _add_accent_bar(d, brand_color)

    role_title = doc.get("role_title", "Job Description")
    title_p = d.add_heading(role_title, level=0)
    for run in title_p.runs:
        run.font.color.rgb = _rgb_color(brand_color)
        run.font.size = Pt(20)

    _add_accent_bar(d, brand_color)

    for heading, body in [
        ("About the Company", doc.get("about_company")),
        ("About the Role",    doc.get("about_role")),
    ]:
        if body:
            _add_heading(d, heading, brand_color)
            d.add_paragraph(body)

    for heading, items in [
        ("Key Responsibilities",    doc.get("responsibilities") or []),
        ("What You'll Need",        doc.get("required_skills") or []),
        ("Nice to Have",            doc.get("preferred_skills") or []),
        ("What Success Looks Like", doc.get("success_criteria") or []),
    ]:
        if items:
            _add_heading(d, heading, brand_color)
            _add_bullet_list(d, items)

    for heading, body in [
        ("Reporting Structure",     doc.get("reporting_structure")),
        ("Growth Path",             doc.get("growth_path")),
        ("Compensation & Benefits", doc.get("compensation")),
    ]:
        if body:
            _add_heading(d, heading, brand_color)
            d.add_paragraph(body)

    eoe = doc.get("equal_opportunity")
    if eoe:
        _add_heading(d, "Equal Opportunity", brand_color)
        p = d.add_paragraph(eoe)
        for run in p.runs:
            run.font.italic = True

    buf = io.BytesIO()
    d.save(buf)
    return buf.getvalue()


def extract_jd_text(doc: JDDocument) -> str:
    """Flatten all prose fields of a JDDocument for the Bedrock OUTPUT guardrail check."""
    parts: list[str] = []
    for key in ("role_title", "about_company", "about_role", "reporting_structure",
                "growth_path", "compensation", "equal_opportunity"):
        val = doc.get(key)
        if val:
            parts.append(val)
    for key in ("responsibilities", "required_skills", "preferred_skills", "success_criteria"):
        items = doc.get(key) or []
        parts.extend(i for i in items if i)
    return "\n\n".join(parts)
